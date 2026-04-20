"""
文件解析模块
支持 PDF、Word (.docx/.doc)、Excel (.xlsx/.xls)、图片 (.jpg/.png/.jpeg) 的解析。
解析结果统一为 dict 结构，供主界面预览和后续审核使用。

PDF 解析优化：优先使用 extract_tables() 提取表格并保留行列结构，
表格外的文本使用 extract_text() 提取。
"""

import base64
import io
import logging
from typing import Any, Dict, List, Optional

import pdfplumber
from pdf2image import convert_from_bytes
from docx import Document as DocxDocument
from PIL import Image
from openpyxl import load_workbook

logger = logging.getLogger(__name__)


# ============================================================
# PDF 解析（优化版：保留表格结构）
# ============================================================
def parse_pdf(file) -> str:
    """解析 PDF 文件，优先提取表格结构，再提取非表格文字。

    当所有页面均无可提取文字时，返回以 ``[扫描件]`` 为前缀的内容，
    供 ``parse_file`` 识别并触发后续 AI-OCR 流程。

    Args:
        file: Streamlit UploadedFile 对象或 file-like 对象。

    Returns:
        提取到的全部文本内容。
    """
    text_parts: list[str] = []
    all_pages_empty: bool = True  # 追踪是否所有页面都没有文字
    try:
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            if len(pdf.pages) == 0:
                return "[提示] 该PDF文件没有任何页面内容。"

            for idx, page in enumerate(pdf.pages, start=1):
                page_texts: list[str] = []
                page_texts.append(f"{'='*20} 第 {idx} 页 {'='*20}")

                # 优先提取表格数据（保留行列结构）
                tables = page.extract_tables()
                table_bboxes = []

                if tables:
                    # 获取表格的边界框用于排除表格区域的文字
                    try:
                        found_tables = page.find_tables()
                        for ft in found_tables:
                            if hasattr(ft, 'bbox'):
                                table_bboxes.append(ft.bbox)
                    except Exception:
                        pass

                    for t_idx, table in enumerate(tables, start=1):
                        page_texts.append(f"\n[表格 {t_idx}]")
                        for row in table:
                            cleaned = [
                                (cell.strip() if cell else "")
                                for cell in row
                            ]
                            page_texts.append("| " + " | ".join(cleaned) + " |")

                # 提取表格区域之外的纯文字内容
                if table_bboxes:
                    # 裁剪掉表格区域后提取文字
                    non_table_text = _extract_text_outside_tables(page, table_bboxes)
                    if non_table_text and non_table_text.strip():
                        # 将非表格文字放在表格之前
                        page_texts.insert(1, non_table_text.strip())
                else:
                    # 没有表格，直接提取全页文字
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        page_texts.append(page_text.strip())

                # 判断本页是否有实际内容（排除页头标题行后长度 > 1 说明有内容）
                if len(page_texts) > 1:
                    all_pages_empty = False

                # 如果整页都没有内容
                if len(page_texts) == 1:
                    page_texts.append("[此页无可提取的文字内容]")

                text_parts.append("\n".join(page_texts))

    except Exception as e:
        logger.error("PDF 解析失败: %s", e)
        return f"[解析失败] 无法解析该PDF文件: {e}"

    full_text = "\n\n".join(text_parts)

    # 如果所有页面都没有文字，标记为扫描件
    if all_pages_empty:
        return "[扫描件] " + full_text

    return full_text


def _extract_text_outside_tables(page, table_bboxes: list) -> str:
    """提取页面中表格区域之外的文字。

    Args:
        page: pdfplumber 页面对象。
        table_bboxes: 表格边界框列表 [(x0, top, x1, bottom), ...]。

    Returns:
        表格区域外的文字。
    """
    try:
        # 获取页面上所有文字及其位置
        words = page.extract_words()
        if not words:
            return ""

        outside_words = []
        for word in words:
            word_center_y = (word.get("top", 0) + word.get("bottom", 0)) / 2
            word_center_x = (word.get("x0", 0) + word.get("x1", 0)) / 2
            in_table = False
            for bbox in table_bboxes:
                x0, top, x1, bottom = bbox
                if x0 <= word_center_x <= x1 and top <= word_center_y <= bottom:
                    in_table = True
                    break
            if not in_table:
                outside_words.append(word.get("text", ""))

        return " ".join(outside_words)
    except Exception:
        # 如果提取失败，回退到全文提取
        try:
            return page.extract_text() or ""
        except Exception:
            return ""


def parse_pdf_structured(file) -> Dict[str, Any]:
    """解析 PDF 文件并返回结构化数据（包含表格的行列结构）。

    用于 report_generator 生成保留原始表格排版的 Excel 报告。

    Args:
        file: Streamlit UploadedFile 对象或 file-like 对象。

    Returns:
        {
            "pages": [
                {
                    "page_num": 1,
                    "text_blocks": ["非表格文字行1", ...],
                    "tables": [
                        [["col1", "col2"], ["val1", "val2"]],  # 二维列表
                    ]
                }
            ]
        }
    """
    result = {"pages": []}
    try:
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            for idx, page in enumerate(pdf.pages, start=1):
                page_data = {
                    "page_num": idx,
                    "text_blocks": [],
                    "tables": [],
                }

                tables = page.extract_tables()
                table_bboxes = []

                if tables:
                    try:
                        found_tables = page.find_tables()
                        for ft in found_tables:
                            if hasattr(ft, 'bbox'):
                                table_bboxes.append(ft.bbox)
                    except Exception:
                        pass

                    for table in tables:
                        cleaned_table = []
                        for row in table:
                            cleaned_row = [
                                (cell.strip() if cell else "") for cell in row
                            ]
                            cleaned_table.append(cleaned_row)
                        page_data["tables"].append(cleaned_table)

                # 非表格文字
                if table_bboxes:
                    non_table_text = _extract_text_outside_tables(page, table_bboxes)
                    if non_table_text and non_table_text.strip():
                        page_data["text_blocks"] = non_table_text.strip().split("\n")
                else:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        page_data["text_blocks"] = page_text.strip().split("\n")

                result["pages"].append(page_data)

    except Exception as e:
        logger.error("PDF 结构化解析失败: %s", e)

    return result


def _pdf_to_images_base64(file, dpi: int = 200, max_pages: int = 10) -> List[str]:
    """将PDF每页转为图片并返回base64编码列表。

    Args:
        file: Streamlit UploadedFile 对象或 file-like 对象。
        dpi: 转换分辨率，默认 200。
        max_pages: 最大处理页数，默认 10。

    Returns:
        每页图片的 base64 编码字符串列表。
    """
    file.seek(0)
    pdf_bytes = file.read()
    file.seek(0)
    images = convert_from_bytes(pdf_bytes, dpi=dpi)
    result: List[str] = []
    for i, img in enumerate(images):
        if i >= max_pages:
            logger.warning("PDF页数超过%d页，仅处理前%d页", max_pages, max_pages)
            break
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        b64 = base64.b64encode(buf.getvalue()).decode("utf-8")

        # 检查图片大小，如果超过 4MB 则降低质量重试
        if len(b64) > 4 * 1024 * 1024:
            logger.warning("第%d页图片过大(%d bytes)，尝试降低分辨率", i + 1, len(b64))
            # 缩小图片尺寸
            max_dim = 2000
            if img.width > max_dim or img.height > max_dim:
                img.thumbnail((max_dim, max_dim), Image.LANCZOS)
            buf = io.BytesIO()
            img.save(buf, format="JPEG", quality=80)
            b64 = base64.b64encode(buf.getvalue()).decode("utf-8")

        result.append(b64)
    return result


# ============================================================
# Word 文档解析
# ============================================================
def parse_docx(file) -> str:
    """解析 Word 文档 (.docx)，提取段落和表格。"""
    filename = getattr(file, "name", "").lower()

    if filename.endswith(".doc") and not filename.endswith(".docx"):
        return _try_parse_old_doc(file)

    try:
        file.seek(0)
        doc = DocxDocument(file)
        parts: list[str] = []

        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            parts.append("--- 文档内容 ---")
            parts.extend(paragraphs)

        if doc.tables:
            for t_idx, table in enumerate(doc.tables, start=1):
                parts.append(f"\n[表格 {t_idx}]")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append("| " + " | ".join(cells) + " |")

        if not parts:
            return "[提示] 该Word文档没有可提取的文字内容。"

        return "\n".join(parts)

    except Exception as e:
        logger.error("DOCX 解析失败: %s", e)
        return f"[解析失败] 无法解析该Word文件: {e}"


def _try_parse_old_doc(file) -> str:
    """尝试解析旧版 .doc 格式。"""
    try:
        file.seek(0)
        doc = DocxDocument(file)
        parts: list[str] = []
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        if paragraphs:
            parts.append("--- 文档内容 ---")
            parts.extend(paragraphs)
        if doc.tables:
            for t_idx, table in enumerate(doc.tables, start=1):
                parts.append(f"\n[表格 {t_idx}]")
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    parts.append("| " + " | ".join(cells) + " |")
        if parts:
            return "\n".join(parts)
        return "[提示] 该Word文档没有可提取的文字内容。"
    except Exception:
        return (
            "[提示] 该文件为旧版 .doc 格式，当前程序无法直接解析。\n"
            "建议您用 Microsoft Word 或 WPS 将文件另存为 .docx 格式后重新上传。"
        )


# ============================================================
# Excel 文档解析
# ============================================================
def parse_xlsx(file) -> str:
    """解析 Excel 文件 (.xlsx/.xls)，遍历所有 sheet 提取内容。"""
    try:
        file.seek(0)
        wb = load_workbook(file, read_only=True, data_only=True)
        parts: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            sheet_parts: list[str] = []
            sheet_parts.append(f"{'='*20} 工作表: {sheet_name} {'='*20}")

            has_content = False
            for row in ws.iter_rows():
                cells = []
                for cell in row:
                    val = cell.value
                    if val is not None:
                        has_content = True
                        cells.append(str(val).strip())
                    else:
                        cells.append("")
                if any(c for c in cells):
                    sheet_parts.append("| " + " | ".join(cells) + " |")

            if not has_content:
                sheet_parts.append("[此工作表无内容]")

            parts.append("\n".join(sheet_parts))

        wb.close()

        if not parts:
            return "[提示] 该Excel文件没有任何工作表内容。"

        return "\n\n".join(parts)

    except Exception as e:
        logger.error("Excel 解析失败: %s", e)
        return f"[解析失败] 无法解析该Excel文件: {e}"


# ============================================================
# 图片处理
# ============================================================
def parse_image(file) -> str:
    """将图片读取并转为 base64 编码字符串。"""
    try:
        file.seek(0)
        raw_bytes = file.read()
        if not raw_bytes:
            return ""
        b64_str = base64.b64encode(raw_bytes).decode("utf-8")
        return b64_str
    except Exception as e:
        logger.error("图片处理失败: %s", e)
        return ""


def get_image_thumbnail(file, max_size: tuple = (300, 300)) -> Optional[Image.Image]:
    """生成图片缩略图供界面预览。"""
    try:
        file.seek(0)
        img = Image.open(file)
        img.thumbnail(max_size)
        return img
    except Exception as e:
        logger.error("缩略图生成失败: %s", e)
        return None


# ============================================================
# 统一入口
# ============================================================
def parse_file(file) -> Dict[str, Any]:
    """根据文件后缀自动选择解析方式。

    Args:
        file: Streamlit UploadedFile 对象。

    Returns:
        {
            "filename": str,
            "type": "pdf" | "docx" | "xlsx" | "image" | "unknown",
            "content": str,          # 提取的文字（图片时为说明文字）
            "is_image": bool,
            "image_base64": str,     # 仅图片类型有值
            "success": bool,         # 解析是否成功
            "structured_data": dict, # PDF 结构化数据（仅 PDF 类型，可选）
        }
    """
    filename: str = getattr(file, "name", "unknown")
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    result: Dict[str, Any] = {
        "filename": filename,
        "type": "unknown",
        "content": "",
        "is_image": False,
        "image_base64": "",
        "success": False,
        "structured_data": None,
    }

    # 空文件检测
    try:
        file.seek(0)
        head = file.read(16)
        file.seek(0)
        if not head:
            result["content"] = "[提示] 该文件为空文件，无内容可解析。"
            return result
    except Exception:
        result["content"] = "[提示] 无法读取文件内容。"
        return result

    # PDF
    if ext == "pdf":
        result["type"] = "pdf"
        content = parse_pdf(file)
        result["content"] = content
        result["success"] = not content.startswith("[解析失败]")

        # 扫描件检测：如果 pdfplumber 提取不到任何文字，标记为扫描件
        if content.startswith("[扫描件]"):
            result["is_scanned_pdf"] = True
            result["success"] = True
            # 将 PDF 每页转为图片 base64，供后续 AI-OCR 使用
            try:
                result["pdf_page_images"] = _pdf_to_images_base64(file)
            except Exception as img_err:
                logger.error("PDF 转图片失败: %s", img_err)
                result["pdf_page_images"] = []
        else:
            result["is_scanned_pdf"] = False
            # 同时获取结构化数据供报告生成使用（仅非扫描件）
            if result["success"]:
                try:
                    result["structured_data"] = parse_pdf_structured(file)
                except Exception:
                    pass

    # Word (.docx / .doc)
    elif ext in ("docx", "doc"):
        result["type"] = "docx"
        content = parse_docx(file)
        result["content"] = content
        result["success"] = not content.startswith("[解析失败]")

    # Excel (.xlsx / .xls)
    elif ext in ("xlsx", "xls"):
        result["type"] = "xlsx"
        if ext == "xls":
            # openpyxl 不支持 .xls 格式，尝试解析，失败则提示
            try:
                content = parse_xlsx(file)
                result["content"] = content
                result["success"] = not content.startswith("[解析失败]")
            except Exception:
                result["content"] = (
                    "[提示] 该文件为旧版 .xls 格式，当前程序无法直接解析。\n"
                    "建议您用 Excel 或 WPS 将文件另存为 .xlsx 格式后重新上传。"
                )
                result["success"] = False
        else:
            content = parse_xlsx(file)
            result["content"] = content
            result["success"] = not content.startswith("[解析失败]")

    # 图片
    elif ext in ("jpg", "jpeg", "png"):
        result["type"] = "image"
        result["is_image"] = True
        b64 = parse_image(file)
        if b64:
            result["image_base64"] = b64
            result["content"] = "图片内容将在审核时由AI识别"
            result["success"] = True
        else:
            result["content"] = "[提示] 图片文件为空或无法读取。"

    # 不支持的格式
    else:
        result["content"] = f"[提示] 不支持的文件格式 (.{ext})，请上传 PDF、Word、Excel 或图片文件。"

    return result
